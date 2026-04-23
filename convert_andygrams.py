"""
Convert Andygram Blog File.xlsx → one .docx per row.

- Filename: {sanitized-title}_ANDYGRAM.docx  (collisions get _2, _3, …)
- Body:     bold title, 'Published: Mon DD, YYYY' line, then prose paragraphs.
- HTML:     <img> stripped; <p>/<br> become paragraph breaks; HTML entities decoded.

Reads: /sessions/tender-happy-carson/mnt/uploads/Andygram Blog File.xlsx
Writes: OUT_DIR (arg 1)  — defaults to afbrain/transcripts/andygrams/
"""
import os, re, sys, html
from datetime import datetime
import openpyxl
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Pt

XLSX = "/sessions/tender-happy-carson/mnt/uploads/Andygram Blog File.xlsx"
DEFAULT_OUT = "/sessions/tender-happy-carson/mnt/afbrain/afbrain/transcripts/andygrams"

# Filenames must be portable; strip anything ugly. Also trim length so the
# _ANDYGRAM suffix always fits under common filesystem limits.
ILLEGAL = re.compile(r'[<>:"/\\|?*\x00-\x1f]')
WHITESPACE = re.compile(r'\s+')

def sanitize_filename(title: str, max_len: int = 120) -> str:
    s = html.unescape(title or "").strip()
    s = ILLEGAL.sub("", s)
    s = WHITESPACE.sub(" ", s)
    s = s.strip(" .")  # Windows dislikes trailing dots/spaces
    if not s:
        s = "Untitled"
    return s[:max_len].rstrip(" .")


def parse_pub_date(raw):
    """Excel stores it as either a datetime or a string like '2026-03-25 07:45:13 -0500'.
    Return a pretty 'Mar 25, 2026' form plus the raw ISO for sorting."""
    if raw is None:
        return "", ""
    if isinstance(raw, datetime):
        dt = raw
    else:
        s = str(raw).strip()
        # Strip trailing timezone offset if present; strptime %z is finicky on
        # some Pythons and we only need the date.
        s = re.sub(r'\s[+-]\d{4}$', '', s)
        for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d", "%m/%d/%Y %H:%M:%S", "%m/%d/%Y"):
            try:
                dt = datetime.strptime(s, fmt)
                break
            except ValueError:
                continue
        else:
            return str(raw), str(raw)
    return dt.strftime("%b %d, %Y"), dt.strftime("%Y-%m-%d")


def html_to_paragraphs(body_html: str) -> list[str]:
    """Turn the post's HTML into a list of plain-text paragraphs.

    Rules:
      - <img> and <figure> are dropped.
      - Each <p> becomes one paragraph.
      - <br><br> inside a <p> splits into multiple paragraphs (Andy uses this a lot).
      - Text outside <p>s is still captured (some posts have bare text/<br>).
      - Entities decoded; collapsed whitespace; empty paragraphs dropped.
    """
    soup = BeautifulSoup(body_html or "", "html.parser")

    # Drop media
    for tag in soup(["img", "figure", "picture", "video", "iframe", "script", "style"]):
        tag.decompose()

    # Normalize: if there are no <p>s at all, wrap everything in one so the
    # loop below still produces output.
    if not soup.find("p"):
        wrapper = soup.new_tag("p")
        for child in list(soup.contents):
            wrapper.append(child.extract())
        soup.append(wrapper)

    paragraphs: list[str] = []
    for p in soup.find_all("p"):
        # Replace each <br> with a sentinel, then split on double-sentinel to
        # honor the <br><br> = paragraph break pattern Andy uses.
        for br in p.find_all("br"):
            br.replace_with("\n")
        text = p.get_text()
        text = html.unescape(text)
        # Split on 2+ newlines, treat single newlines as soft breaks inside a
        # single paragraph (preserved as spaces since "normal paragraphs").
        chunks = re.split(r'\n\s*\n+', text)
        for chunk in chunks:
            cleaned = re.sub(r'\s+', ' ', chunk).strip()
            if cleaned:
                paragraphs.append(cleaned)
    return paragraphs


def build_docx(title: str, pub_pretty: str, paragraphs: list[str], out_path: str) -> None:
    doc = Document()
    # Base font → Calibri 12 (Word default). Title is bold + larger.
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(12)

    t = doc.add_paragraph()
    run = t.add_run(title)
    run.bold = True
    run.font.size = Pt(18)

    if pub_pretty:
        d = doc.add_paragraph()
        dr = d.add_run(f"Published: {pub_pretty}")
        dr.italic = True
        dr.font.size = Pt(10)

    # Spacer
    doc.add_paragraph("")

    for para in paragraphs:
        doc.add_paragraph(para)

    doc.save(out_path)


def main():
    out_dir = sys.argv[1] if len(sys.argv) > 1 else DEFAULT_OUT
    os.makedirs(out_dir, exist_ok=True)
    print(f"Output dir: {out_dir}")

    wb = openpyxl.load_workbook(XLSX, data_only=True)
    ws = wb.active

    # Track seen filenames to disambiguate duplicate titles.
    used: dict[str, int] = {}
    written = 0
    skipped_no_body = 0
    errors: list[tuple[int, str]] = []

    for i, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
        _id, title, body_html, pub_at, _handle = row
        if not body_html or not str(body_html).strip():
            skipped_no_body += 1
            continue

        title = (title or f"Andygram_{_id}").strip()
        base = sanitize_filename(title)
        # Collision handling: append _2, _3, …
        key = base.lower()
        used[key] = used.get(key, 0) + 1
        suffix = "" if used[key] == 1 else f"_{used[key]}"
        filename = f"{base}{suffix}_ANDYGRAM.docx"
        out_path = os.path.join(out_dir, filename)

        pub_pretty, _ = parse_pub_date(pub_at)

        try:
            paragraphs = html_to_paragraphs(body_html)
            if not paragraphs:
                # Fallback: raw text, in case HTML parse yielded nothing.
                raw = re.sub(r'<[^>]+>', ' ', str(body_html))
                raw = html.unescape(raw)
                raw = re.sub(r'\s+', ' ', raw).strip()
                if raw:
                    paragraphs = [raw]
            if not paragraphs:
                skipped_no_body += 1
                continue
            build_docx(title, pub_pretty, paragraphs, out_path)
            written += 1
        except Exception as e:
            errors.append((i, f"{title!r}: {e}"))

        if written % 250 == 0 and written:
            print(f"  … wrote {written}")

    print(f"\nDone. Wrote {written} docs. Skipped {skipped_no_body} (no body). Errors: {len(errors)}")
    for lineno, msg in errors[:10]:
        print(f"  row {lineno}: {msg}")


if __name__ == "__main__":
    main()
